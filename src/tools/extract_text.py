import sys
from pathlib import Path


def extract_pdf(file_path: Path) -> str:
    import fitz

    doc = fitz.open(str(file_path))
    text_parts: list[str] = []
    for page_idx in range(len(doc)):
        page = doc[page_idx]
        text = page.get_text("text")
        if text.strip():
            text_parts.append(f"--- Page {page_idx + 1} ---\n{text}")
    doc.close()
    return "\n\n".join(text_parts)


def extract_docx(file_path: Path) -> str:
    from docx import Document

    doc = Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables_text: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            tables_text.append(" | ".join(cells))
    result = "\n\n".join(paragraphs)
    if tables_text:
        result += "\n\n--- Tables ---\n" + "\n".join(tables_text)
    return result


def extract_markdown(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8")


def extract_text(file_path: str) -> str:
    path = Path(file_path)
    if not path.exists():
        print(f"Error: file not found: {file_path}", file=sys.stderr)
        sys.exit(1)

    suffix = path.suffix.lower()
    extractors = {
        ".pdf": extract_pdf,
        ".docx": extract_docx,
        ".doc": extract_docx,
        ".md": extract_markdown,
        ".txt": extract_markdown,
        ".markdown": extract_markdown,
    }

    extractor = extractors.get(suffix)
    if extractor is None:
        print(f"Error: unsupported format: {suffix}", file=sys.stderr)
        print(f"Supported: {', '.join(extractors.keys())}", file=sys.stderr)
        sys.exit(1)

    return extractor(path)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python -m src.tools.extract_text <file>", file=sys.stderr)
        sys.exit(1)
    print(extract_text(sys.argv[1]))
